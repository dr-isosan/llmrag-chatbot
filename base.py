import os
import json
import re
import hashlib
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging

# Document processing libraries
import fitz  # PyMuPDF
from docx import Document

# textract not available - using alternative approach
from config import config

logger = logging.getLogger(__name__)


@dataclass
class DocumentMetadata:
    """Doküman metadata'sı için dataclass"""

    filename: str
    file_type: str
    file_size: int
    creation_date: Optional[str] = None
    modification_date: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    page_count: Optional[int] = None
    language: Optional[str] = None
    checksum: Optional[str] = None


@dataclass
class ProcessedChunk:
    """İşlenmiş chunk için dataclass"""

    content: str
    chunk_index: int
    start_pos: int
    end_pos: int
    chunk_type: str = "content"  # content, header, table, etc.
    metadata: Optional[Dict[str, Any]] = None


class AdvancedDocumentProcessor:
    """Gelişmiş doküman işleme sınıfı"""

    SUPPORTED_FORMATS = {
        ".pdf": "PDF",
        ".docx": "DOCX",
        ".doc": "DOC",
        ".txt": "TXT",
        ".rtf": "RTF",
        ".html": "HTML",
        ".htm": "HTML",
        ".md": "MARKDOWN",
    }

    def __init__(self):
        self.chunk_size = config.MAX_CHUNK_SIZE
        self.overlap = config.CHUNK_OVERLAP
        self.stats = {
            "processed_files": 0,
            "failed_files": 0,
            "total_chunks": 0,
            "total_characters": 0,
        }

    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Gelişmiş PDF metin çıkarma"""
        try:
            doc = fitz.open(file_path)
            text_parts = []
            metadata = self._extract_pdf_metadata(doc, file_path)

            for page_num in range(len(doc)):
                page = doc[page_num]
                # Metin çıkarma
                text = page.get_text()  # type: ignore

                # Tablo tespiti ve çıkarma
                tables = page.find_tables() if hasattr(page, "find_tables") else None  # type: ignore
                if tables:
                    for table in tables:
                        try:
                            table_data = (
                                table.extract() if hasattr(table, "extract") else None
                            )
                            table_text = (
                                self._format_table_text(table_data)
                                if table_data
                                else ""
                            )
                            text += f"\n[TABLO {page_num}]\n{table_text}\n[/TABLO]\n"
                        except Exception as e:
                            logger.warning(
                                f"Tablo çıkarma hatası sayfa {page_num}: {e}"
                            )

                # Görsel-metin ilişkisi (OCR için placeholder)
                images = page.get_images() if hasattr(page, "get_images") else None  # type: ignore
                if images:
                    text += f"\n[GÖRSEL SAYISI: {len(images)}]\n"

                if text.strip():
                    text_parts.append(
                        f"[SAYFA {page_num}]\n{text}\n[/SAYFA {page_num}]"
                    )
                else:
                    logger.debug(f"Sayfa {page_num} boş")

            doc.close()
            full_text = "\n".join(text_parts)
            metadata.page_count = len(text_parts)

            logger.info(
                f"   📄 PDF'den {len(text_parts)} sayfa metin çıkarıldı, toplam {len(full_text)} karakter"
            )

            return full_text, metadata

        except Exception as e:
            logger.error(f"PDF okuma hatası {file_path}: {e}")
            return "", self._create_error_metadata(file_path, str(e))

    def extract_text_from_docx(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Gelişmiş DOCX metin çıkarma"""
        try:
            doc = Document(file_path)
            text_parts = []
            metadata = self._extract_docx_metadata(doc, file_path)

            # Paragraflar
            for para in doc.paragraphs:
                if para.text.strip():
                    # Başlık tespiti
                    if (
                        para.style
                        and getattr(para.style, "name", None)
                        and str(para.style.name).startswith("Heading")
                    ):
                        heading_name = str(para.style.name)
                        heading_level = (
                            heading_name[-1]
                            if heading_name and heading_name[-1].isdigit()
                            else ""
                        )
                        text_parts.append(
                            f"[BAŞLIK {heading_level}] {para.text} [/BAŞLIK]"
                        )
                    else:
                        text_parts.append(para.text)

            # Tablolar
            for table_num, table in enumerate(doc.tables, 1):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                table_text = self._format_table_text(table_data)
                text_parts.append(f"\n[TABLO {table_num}]\n{table_text}\n[/TABLO]\n")

            return "\n".join(text_parts), metadata

        except Exception as e:
            logger.error(f"DOCX okuma hatası {file_path}: {e}")
            return "", self._create_error_metadata(file_path, str(e))

    def extract_text_universal(self, file_path: str) -> Tuple[str, DocumentMetadata]:
        """Diğer formatlar için universal çıkarma (textract)"""
        try:
            try:
                import textract  # type: ignore
            except ImportError:
                logger.error(
                    "textract kütüphanesi yüklü değil. Universal extraction kullanılamaz."
                )
                return "", self._create_error_metadata(
                    file_path, "textract not installed"
                )
            text = textract.process(file_path).decode("utf-8")
            metadata = self._create_basic_metadata(file_path)
            return text, metadata
        except Exception as e:
            logger.error(f"Universal text extraction hatası {file_path}: {e}")
            return "", self._create_error_metadata(file_path, str(e))

    def _extract_pdf_metadata(
        self, doc: fitz.Document, file_path: str
    ) -> DocumentMetadata:
        """PDF metadata çıkarma"""
        try:
            pdf_metadata = getattr(doc, "metadata", None)
            file_stats = os.stat(file_path)

            return DocumentMetadata(
                filename=os.path.basename(file_path),
                file_type="PDF",
                file_size=file_stats.st_size,
                creation_date=(
                    pdf_metadata.get("creationDate", "") if pdf_metadata else None
                ),
                modification_date=(
                    pdf_metadata.get("modDate", "") if pdf_metadata else None
                ),
                author=pdf_metadata.get("author", "") if pdf_metadata else None,
                title=pdf_metadata.get("title", "") if pdf_metadata else None,
                page_count=getattr(doc, "page_count", None),
                checksum=self._calculate_checksum(file_path),
            )
        except Exception as e:
            logger.warning(f"PDF metadata çıkarma hatası: {e}")
            return self._create_basic_metadata(file_path)

    def _extract_docx_metadata(self, doc, file_path: str) -> DocumentMetadata:
        """DOCX metadata çıkarma"""
        try:
            core_props = doc.core_properties
            file_stats = os.stat(file_path)

            return DocumentMetadata(
                filename=os.path.basename(file_path),
                file_type="DOCX",
                file_size=file_stats.st_size,
                creation_date=str(core_props.created) if core_props.created else None,
                modification_date=(
                    str(core_props.modified) if core_props.modified else None
                ),
                author=core_props.author,
                title=core_props.title,
                checksum=self._calculate_checksum(file_path),
            )
        except Exception as e:
            logger.warning(f"DOCX metadata çıkarma hatası: {e}")
            return self._create_basic_metadata(file_path)

    def _create_basic_metadata(self, file_path: str) -> DocumentMetadata:
        """Temel metadata oluşturma"""
        file_stats = os.stat(file_path)
        file_ext = Path(file_path).suffix.lower()

        return DocumentMetadata(
            filename=os.path.basename(file_path),
            file_type=self.SUPPORTED_FORMATS.get(file_ext, "UNKNOWN"),
            file_size=file_stats.st_size,
            checksum=self._calculate_checksum(file_path),
        )

    def _create_error_metadata(self, file_path: str, error: str) -> DocumentMetadata:
        """Hata durumu için metadata"""
        metadata = self._create_basic_metadata(file_path)
        # 'metadata' alanı yok, ek bilgi için checksum'a hata ekle
        metadata.checksum = f"error:{error}"
        return metadata

    def _calculate_checksum(self, file_path: str) -> str:
        """Dosya checksum hesaplama"""
        try:
            hasher = hashlib.md5()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hasher.update(chunk)
            return hasher.hexdigest()
        except Exception as e:
            logger.warning(f"Checksum hesaplama hatası: {e}")
            return ""

    def _format_table_text(self, table_data: List[List[str]]) -> str:
        """Tablo verilerini metin formatına çevirme"""
        if not table_data:
            return ""

        formatted_rows = []
        for row in table_data:
            # None hücreleri boş stringe çevir
            clean_row = [
                str(cell).strip() if cell is not None else ""
                for cell in row
                if cell is not None and str(cell).strip()
            ]
            if clean_row:
                formatted_rows.append(" | ".join(clean_row))

        return "\n".join(formatted_rows)

    def advanced_clean_text(self, text: str, preserve_structure: bool = True) -> str:
        """Gelişmiş metin temizleme"""
        if not text:
            return ""

        original_text = text

        # Türkçe-spesifik düzeltmeler
        turkish_corrections = {
            "İ": "i",
            "I": "ı",
            "Ğ": "ğ",
            "Ü": "ü",
            "Ş": "ş",
            "Ö": "ö",
            "Ç": "ç",
        }

        # Yapısal etiketleri geçici olarak koru
        structure_markers = re.findall(r"\[(?:SAYFA|BAŞLIK|TABLO|GÖRSEL)[^\]]*\]", text)

        if preserve_structure:
            # Yapısal bilgileri koru
            pass
        else:
            # Yapısal etiketleri kaldır
            text = re.sub(r"\[(?:SAYFA|BAŞLIK|TABLO|GÖRSEL)[^\]]*\]", "", text)

        # Temel temizlik
        # Çoklu boşlukları tek boşluk yap
        text = re.sub(r"\s+", " ", text)

        # Özel karakterleri temizle (Türkçe karakterleri koru)
        text = re.sub(r'[^\w\sçÇğĞıİöÖşŞüÜ.,;:!?\-\(\)\[\]"\'\/]', "", text)

        # Gereksiz sayfa numaralarını temizle (ama önemli sayıları koru)
        text = re.sub(r"\bSayfa\s*\d+\b", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\bPage\s*\d+\b", "", text, flags=re.IGNORECASE)

        # Başlık numaralarını kontrollü temizle
        # Sadece satır başındaki numaralandırmayı temizle
        text = re.sub(r"^\d+(\.\d+)*\s+", "", text, flags=re.MULTILINE)

        # Çoklu satır sonlarını normalize et
        text = re.sub(r"\n\s*\n+", "\n\n", text)

        # Başındaki ve sonundaki boşlukları temizle
        text = text.strip()

        # Sonuç kontrolü
        if len(text) < len(original_text) * 0.1:  # %90'dan fazla kayıp varsa
            logger.warning("Aşırı temizlik algılandı, orijinal metin korunuyor")
            return original_text.strip()

        return text

    def adaptive_chunk_creation(
        self, text: str, metadata: DocumentMetadata
    ) -> List[ProcessedChunk]:
        """Adaptif chunk oluşturma"""
        if not text:
            return []

        chunks = []

        # Doküman tipine göre stratejik chunking
        if metadata.file_type == "PDF" and metadata.page_count:
            # Sayfa bazlı chunking
            chunks.extend(self._page_based_chunking(text))
        else:
            # Paragraf bazlı chunking
            chunks.extend(self._paragraph_based_chunking(text))

        # Chunk kalite kontrolü
        chunks = self._validate_chunks(chunks)

        return chunks

    def _page_based_chunking(self, text: str) -> List[ProcessedChunk]:
        """Sayfa bazlı chunking"""
        chunks = []
        page_pattern = r"\[SAYFA (\d+)\](.*?)\[/SAYFA \1\]"

        page_matches = re.finditer(page_pattern, text, re.DOTALL)
        page_matches_list = list(page_matches)

        logger.info(
            f"   🔍 Sayfa bazlı chunking: {len(page_matches_list)} sayfa bulundu"
        )

        if not page_matches_list:
            logger.warning(
                f"   ⚠️ Sayfa etiketleri bulunamadı, paragraf chunking'e geçiliyor"
            )
            return self._paragraph_based_chunking(text)

        for match in page_matches_list:
            page_num = match.group(1)
            page_content = match.group(2).strip()

            if len(page_content) <= self.chunk_size:
                # Sayfa küçükse tek chunk
                chunks.append(
                    ProcessedChunk(
                        content=page_content,
                        chunk_index=len(chunks),
                        start_pos=match.start(),
                        end_pos=match.end(),
                        chunk_type="page",
                        metadata={"page_number": page_num},
                    )
                )
            else:
                # Büyük sayfaları böl
                sub_chunks = self._split_large_content(page_content)
                for i, sub_chunk in enumerate(sub_chunks):
                    chunks.append(
                        ProcessedChunk(
                            content=sub_chunk,
                            chunk_index=len(chunks),
                            start_pos=match.start(),
                            end_pos=match.end(),
                            chunk_type="page_part",
                            metadata={"page_number": page_num, "part": i + 1},
                        )
                    )

        return chunks

    def _paragraph_based_chunking(self, text: str) -> List[ProcessedChunk]:
        """Paragraf bazlı chunking"""
        chunks = []
        paragraphs = text.split("\n\n")

        current_chunk = ""
        current_start = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Eğer mevcut chunk + yeni paragraf çok büyükse
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                # Mevcut chunk'ı kaydet
                chunks.append(
                    ProcessedChunk(
                        content=current_chunk.strip(),
                        chunk_index=len(chunks),
                        start_pos=current_start,
                        end_pos=current_start + len(current_chunk),
                        chunk_type="paragraph_group",
                    )
                )

                # Yeni chunk başlat (overlap ile)
                current_chunk = current_chunk[-self.overlap :] + " " + para
                current_start = current_start + len(current_chunk) - self.overlap
            else:
                # Paragrafi ekle
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para

        # Son chunk'ı ekle
        if current_chunk.strip():
            chunks.append(
                ProcessedChunk(
                    content=current_chunk.strip(),
                    chunk_index=len(chunks),
                    start_pos=current_start,
                    end_pos=current_start + len(current_chunk),
                    chunk_type="paragraph_group",
                )
            )

        return chunks

    def _split_large_content(self, content: str) -> List[str]:
        """Büyük içeriği böl"""
        if len(content) <= self.chunk_size:
            return [content]

        chunks = []
        words = content.split()
        current_chunk_words = []
        current_length = 0

        for word in words:
            word_length = len(word) + 1  # +1 for space

            if current_length + word_length > self.chunk_size and current_chunk_words:
                # Chunk'ı kaydet
                chunks.append(" ".join(current_chunk_words))

                # Overlap için son birkaç kelimeyi al
                overlap_words = current_chunk_words[
                    -self.overlap // 10 :
                ]  # Approximate overlap
                current_chunk_words = overlap_words + [word]
                current_length = sum(len(w) + 1 for w in current_chunk_words)
            else:
                current_chunk_words.append(word)
                current_length += word_length

        # Son chunk'ı ekle
        if current_chunk_words:
            chunks.append(" ".join(current_chunk_words))

        return chunks

    def _validate_chunks(self, chunks: List[ProcessedChunk]) -> List[ProcessedChunk]:
        """Chunk kalite kontrolü"""
        validated_chunks = []

        for chunk in chunks:
            # Çok kısa chunk'ları filtrele
            if len(chunk.content.strip()) < 20:
                logger.debug(f"Çok kısa chunk atlandı: {len(chunk.content)} karakter")
                continue

            # Sadece noktalama işareti içeren chunk'ları filtrele
            if re.match(r"^[^\w]*$", chunk.content.strip()):
                logger.debug("Sadece noktalama işareti içeren chunk atlandı")
                continue

            validated_chunks.append(chunk)

        return validated_chunks

    def process_documents(self, path: str, keyword: str = None) -> List[Dict[str, Any]]:
        """Ana doküman işleme fonksiyonu (anahtar kelime eşleşmesi zorunlu)"""
        if not os.path.exists(path):
            raise FileNotFoundError(f"Dosya veya klasör bulunamadı: {path}")

        # Dosya listesi oluştur
        files_to_process = self._get_supported_files(path)

        if not files_to_process:
            logger.warning(f"İşlenebilir dosya bulunamadı: {path}")
            return []

        logger.info(f"📄 {len(files_to_process)} dosya bulundu. İşleniyor...")

        processed_data = []

        for i, file_path in enumerate(files_to_process, 1):
            logger.info(
                f"[{i}/{len(files_to_process)}] İşleniyor: {os.path.basename(file_path)}"
            )

            try:
                # Dosya tipine göre işle
                file_ext = Path(file_path).suffix.lower()

                if file_ext == ".pdf":
                    raw_text, metadata = self.extract_text_from_pdf(file_path)
                elif file_ext in [".docx", ".doc"]:
                    raw_text, metadata = self.extract_text_from_docx(file_path)
                else:
                    raw_text, metadata = self.extract_text_universal(file_path)

                if not raw_text.strip():
                    logger.warning(f"   ⚠️ Boş içerik: {metadata.filename}")
                    self.stats["failed_files"] += 1
                    continue

                # Metni temizle
                cleaned_text = self.advanced_clean_text(raw_text)

                logger.info(
                    f"   🧹 Temizleme: {len(raw_text)} -> {len(cleaned_text)} karakter"
                )

                if not cleaned_text.strip():
                    logger.warning(
                        f"   ⚠️ Temizleme sonrası boş içerik: {metadata.filename}"
                    )
                    self.stats["failed_files"] += 1
                    continue

                # Adaptif chunking
                processed_chunks = self.adaptive_chunk_creation(cleaned_text, metadata)

                logger.info(
                    f"   📦 Chunking sonucu: {len(processed_chunks)} chunk oluşturuldu"
                )

                if not processed_chunks:
                    logger.warning(f"   ⚠️ Chunk oluşturulamadı: {metadata.filename}")
                    self.stats["failed_files"] += 1
                    continue

                # Sonuç verilerini hazırla
                chunk_texts = [chunk.content for chunk in processed_chunks]
                chunk_metadata = [chunk.metadata or {} for chunk in processed_chunks]

                result_data = {
                    "filename": metadata.filename,
                    "file_type": metadata.file_type,
                    "file_size": metadata.file_size,
                    "content": cleaned_text,
                    "chunks": chunk_texts,
                    "chunk_metadata": chunk_metadata,
                    "chunk_count": len(chunk_texts),
                    "character_count": len(cleaned_text),
                    "document_metadata": {
                        "creation_date": metadata.creation_date,
                        "modification_date": metadata.modification_date,
                        "author": metadata.author,
                        "title": metadata.title,
                        "page_count": metadata.page_count,
                        "checksum": metadata.checksum,
                    },
                }
                # Anahtar kelimeyi sadece bu dosya için ekle
                if keyword is not None and os.path.basename(file_path) == result_data["filename"]:
                    result_data["keyword"] = keyword

                processed_data.append(result_data)

                # İstatistik güncelle
                self.stats["processed_files"] += 1
                self.stats["total_chunks"] += len(chunk_texts)
                self.stats["total_characters"] += len(cleaned_text)

                logger.info(
                    f"   ✅ {len(chunk_texts)} chunk, {len(cleaned_text):,} karakter"
                )

            except Exception as e:
                logger.error(f"   ❌ İşleme hatası {os.path.basename(file_path)}: {e}")
                self.stats["failed_files"] += 1
                continue

        # Final istatistikler
        logger.info(f"📊 İşlem tamamlandı:")
        logger.info(f"   Başarılı: {self.stats['processed_files']}")
        logger.info(f"   Başarısız: {self.stats['failed_files']}")
        logger.info(f"   Toplam chunk: {self.stats['total_chunks']:,}")
        logger.info(f"   Toplam karakter: {self.stats['total_characters']:,}")

        return processed_data

    def _get_supported_files(self, path: str) -> List[str]:
        """Desteklenen dosyaları listele"""
        files = []

        if os.path.isfile(path):
            if Path(path).suffix.lower() in self.SUPPORTED_FORMATS:
                files.append(path)
        elif os.path.isdir(path):
            for root, _, filenames in os.walk(path):
                for filename in filenames:
                    if Path(filename).suffix.lower() in self.SUPPORTED_FORMATS:
                        files.append(os.path.join(root, filename))

        return sorted(files)  # Deterministic order


def save_enhanced_data(data: List[Dict[str, Any]], output_file: str) -> None:
    """Gelişmiş veri kaydetme"""
    try:
        # Backup eski dosya
        if os.path.exists(output_file):
            backup_file = f"{output_file}.backup"
            os.rename(output_file, backup_file)
            logger.info(f"Eski dosya yedeklendi: {backup_file}")

        # Yeni dosyayı kaydet
        with open(output_file, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

        logger.info(f"✅ Veriler kaydedildi: {output_file}")

        # Dosya boyutu kontrolü
        file_size = os.path.getsize(output_file) / (1024 * 1024)  # MB
        logger.info(f"📁 Dosya boyutu: {file_size:.2f} MB")

    except Exception as e:
        logger.error(f"❌ Kaydetme hatası: {e}")
        raise


def main():
    """uploads klasöründeki tüm desteklenen dosyaları işleyip uploads_base.json'a kaydeder"""
    uploads_dir = "uploads"
    output_file = "uploads_base.json"

    # Logging setup
    logging.basicConfig(
        level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s"
    )

    try:
        logger.info("🚀 uploads klasöründeki dosyalar işleniyor...")

        processor = AdvancedDocumentProcessor()

        if not os.path.exists(uploads_dir):
            logger.error(f"uploads klasörü bulunamadı: {uploads_dir}")
            return

        files_to_process = processor._get_supported_files(uploads_dir)

        if not files_to_process:
            logger.error("❌ uploads klasöründe işlenebilir dosya bulunamadı.")
            return

        all_data = []
        for i, file_path in enumerate(files_to_process, 1):
            logger.info(f"[{i}/{len(files_to_process)}] İşleniyor: {file_path}")
            try:
                data = processor.process_documents(file_path)
                if data:
                    all_data.extend(data)
                    logger.info(f"✅ {file_path} işlendi ve eklendi.")
                else:
                    logger.warning(f"⚠️ İşlenemeyen dosya: {file_path}")
            except Exception as e:
                logger.error(f"❌ Dosya işleme hatası: {file_path}, Hata: {e}")

        if all_data:
            save_enhanced_data(all_data, output_file)
            logger.info(f"✅ Tüm dosyalar {output_file} dosyasına kaydedildi.")
        else:
            logger.warning("⚠️ Hiçbir dosya başarıyla işlenemedi.")

    except Exception as e:
        logger.error(f"❌ Ana işlem hatası: {e}")
        raise


if __name__ == "__main__":
    main()